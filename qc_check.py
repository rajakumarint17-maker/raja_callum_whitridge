import sys, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

base = 'C:/Users/HP/Downloads/kensei/openclaw-intelligence/task_folder/data/'
targets = ['doc_44.docx', 'doc_18.docx']

dollar_pattern = re.compile(r'\$[\d,\.]+')
ratio_pattern = re.compile(r'(1\s*(per|:|to)\s*8|one\s+per\s+eight|chaperone.{0,30}ratio|ratio.{0,30}chaperone|1:8)', re.IGNORECASE)
keywords = ['chaperone', 'ratio', 'supervisor', '1 per', 'per 8', '863', '243', '9.00', '620', '900', '200']

for fname in targets:
    doc = Document(base + fname)
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)
    full = '\n'.join(all_text)
    print('=== ' + fname + ' SEARCH RESULTS ===')
    dollar_hits = dollar_pattern.findall(full)
    print('Dollar values found: ' + str(dollar_hits))
    ratio_hits = ratio_pattern.findall(full)
    print('Ratio patterns found: ' + str(ratio_hits))
    for kw in keywords:
        hits = [line for line in all_text if kw.lower() in line.lower()]
        if hits:
            print('Keyword [' + kw + ']: ' + str(hits))
    print()
